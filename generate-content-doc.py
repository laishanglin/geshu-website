from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# ── 样式设置 ──
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)

# ── 标题 ──
title = doc.add_heading('格数科技官网 — 内容清单', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    '📋 说明：请在此文档中直接修改需要更新的文字内容。修改完成后将文档发回给我，我会根据修改后的内容更新网站。\n'
    '📌 请不要修改结构/表格格式，只修改【内容】列中的文字。',
    style='Normal'
).italic = True

doc.add_paragraph()  # 空行

# ════════════════════════════════════════════════════════
# 1. 全局站点信息
# ════════════════════════════════════════════════════════
doc.add_heading('一、全局站点信息', level=1)
doc.add_paragraph('数据文件：src/data/site.json', style='Normal').italic = True

table1 = doc.add_table(rows=9, cols=3)
table1.style = 'Light Grid Accent 1'

headers1 = ['字段', '当前内容', '修改后内容（留空=不改）']
for i, h in enumerate(headers1):
    cell = table1.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].font.bold = True

site_data = [
    ('公司名（中文）', 'companyName', '格数科技'),
    ('公司名（英文）', 'companyNameEn', 'Geshu Technology'),
    ('Slogan', 'slogan', '让数据成为企业的核心驱动力'),
    ('描述', 'description', '格数科技——企业级数据智能解决方案提供商。我们致力于帮助企业挖掘数据价值，实现数字化转型升级。'),
    ('成立年份', 'foundedYear', '2025'),
    ('联系电话', 'contactPhone', '（空）'),
    ('联系邮箱', 'contactEmail', 'info@geshu-tech.com'),
    ('联系地址', 'contactAddress', '（空）'),
]

for i, (label, key, value) in enumerate(site_data):
    table1.rows[i+1].cells[0].text = label
    table1.rows[i+1].cells[1].text = value
    table1.rows[i+1].cells[2].text = ''

doc.add_paragraph()

# ════════════════════════════════════════════════════════
# 2. 导航栏
# ════════════════════════════════════════════════════════
doc.add_heading('二、导航栏', level=1)
doc.add_paragraph('数据文件：src/data/navigation.json', style='Normal').italic = True

table2 = doc.add_table(rows=8, cols=3)
table2.style = 'Light Grid Accent 1'

headers2 = ['位置', '当前内容', '修改后内容（留空=不改）']
for i, h in enumerate(headers2):
    table2.rows[0].cells[i].text = h
    table2.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

nav_items = [
    ('导航链接 1', '首页'),
    ('导航链接 2', '关于我们'),
    ('导航链接 3', '核心业务'),
    ('导航链接 4', '使命愿景'),
    ('导航链接 5', '新闻动态'),
    ('导航链接 6', '联系我们'),
    ('CTA按钮（导航栏右侧）', '立即咨询'),
]

for i, (label, value) in enumerate(nav_items):
    table2.rows[i+1].cells[0].text = label
    table2.rows[i+1].cells[1].text = value
    table2.rows[i+1].cells[2].text = ''

doc.add_paragraph()

# ════════════════════════════════════════════════════════
# 3. Hero 区域
# ════════════════════════════════════════════════════════
doc.add_heading('三、Hero 首屏区域', level=1)
doc.add_paragraph('数据文件：src/data/hero.json', style='Normal').italic = True

table3 = doc.add_table(rows=5, cols=3)
table3.style = 'Light Grid Accent 1'

for i, h in enumerate(['字段', '当前内容', '修改后内容（留空=不改）']):
    table3.rows[0].cells[i].text = h
    table3.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

hero_items = [
    ('Slogan 前半句', '让数据成为企业的'),
    ('Slogan 高亮后半句', '核心驱动力'),
    ('副标题', '格数科技——企业级数据智能解决方案提供商'),
    ('CTA按钮 1', '了解核心业务'),
]

for i, (label, value) in enumerate(hero_items):
    table3.rows[i+1].cells[0].text = label
    table3.rows[i+1].cells[1].text = value
    table3.rows[i+1].cells[2].text = ''

doc.add_paragraph()

# ════════════════════════════════════════════════════════
# 4. 关于我们
# ════════════════════════════════════════════════════════
doc.add_heading('四、「关于我们」区域', level=1)
doc.add_paragraph('数据文件：src/data/company.json + 组件内硬编码（CompanyOverview.jsx）', style='Normal').italic = True

doc.add_heading('板块标题 & 简介', level=2)
table4a = doc.add_table(rows=4, cols=3)
table4a.style = 'Light Grid Accent 1'
for i, h in enumerate(['字段', '当前内容', '修改后内容（留空=不改）']):
    table4a.rows[0].cells[i].text = h
    table4a.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

company_title_items = [
    ('板块中文标题', '关于格数科技'),
    ('板块英文标题', 'ABOUT US'),
    ('简介文字', '格数科技成立于2025年，是一家专注于企业数据智能化的科技公司。我们坚信数据是数字经济时代最核心的生产要素，致力于帮助企业构建数据驱动的决策体系和智能化运营能力。'),
]
for i, (label, value) in enumerate(company_title_items):
    table4a.rows[i+1].cells[0].text = label
    table4a.rows[i+1].cells[1].text = value
    table4a.rows[i+1].cells[2].text = ''

doc.add_paragraph()

doc.add_heading('三大亮点卡片', level=2)
table4b = doc.add_table(rows=4, cols=4)
table4b.style = 'Light Grid Accent 1'
for i, h in enumerate(['卡片', '标题', '描述', '修改后内容（直接改写整行）']):
    table4b.rows[0].cells[i].text = h
    table4b.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

highlights = [
    ('1', '技术驱动', '自主研发数据智能平台，融合大数据、AI与行业Know-how'),
    ('2', '专家团队', '核心团队来自头部科技企业与咨询公司，兼具技术深度与业务洞察'),
    ('3', '端到端服务', '从战略规划到落地实施，提供全链路数据解决方案'),
]
for i, (num, title, desc) in enumerate(highlights):
    table4b.rows[i+1].cells[0].text = num
    table4b.rows[i+1].cells[1].text = title
    table4b.rows[i+1].cells[2].text = desc
    table4b.rows[i+1].cells[3].text = ''

doc.add_paragraph()

doc.add_heading('右侧堆叠卡片（硬编码在 CompanyOverview.jsx 中）', level=2)
table4c = doc.add_table(rows=5, cols=4)
table4c.style = 'Light Grid Accent 1'
for i, h in enumerate(['卡片', '标题', '描述', '修改后内容（直接改写整行）']):
    table4c.rows[0].cells[i].text = h
    table4c.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

stack_cards = [
    ('1', '数据资产', '打通数据孤岛，构建企业级数据底座，让数据真正可用、可控、可运营。'),
    ('2', '智能分析', '融合AI与行业Know-how，让数据洞察驱动精准决策，释放企业增长潜能。'),
    ('3', '业务增长', '从数据到洞察，从洞察到行动，从行动到增长——完成智能化的完整闭环。'),
    ('底部标注', '（visualDesc）', '数据驱动的企业智能化转型'),
]
for i, (num, title, desc) in enumerate(stack_cards):
    table4c.rows[i+1].cells[0].text = num
    table4c.rows[i+1].cells[1].text = title
    table4c.rows[i+1].cells[2].text = desc
    table4c.rows[i+1].cells[3].text = ''

doc.add_paragraph()

# ════════════════════════════════════════════════════════
# 5. 核心业务
# ════════════════════════════════════════════════════════
doc.add_heading('五、「核心业务」区域', level=1)
doc.add_paragraph('数据文件：src/data/businessAreas.json', style='Normal').italic = True

doc.add_heading('板块标题', level=2)
table5a = doc.add_table(rows=4, cols=3)
table5a.style = 'Light Grid Accent 1'
for i, h in enumerate(['字段', '当前内容', '修改后内容（留空=不改）']):
    table5a.rows[0].cells[i].text = h
    table5a.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

ba_title_items = [
    ('板块中文标题', '核心业务'),
    ('板块英文标题', 'WHAT WE DO'),
    ('副标题', '四大业务方向，构建企业数据智能全链路'),
]
for i, (label, value) in enumerate(ba_title_items):
    table5a.rows[i+1].cells[0].text = label
    table5a.rows[i+1].cells[1].text = value
    table5a.rows[i+1].cells[2].text = ''

doc.add_paragraph()

doc.add_heading('业务卡片', level=2)
table5b = doc.add_table(rows=5, cols=4)
table5b.style = 'Light Grid Accent 1'
for i, h in enumerate(['卡片ID', '标题 + 描述 + 标签', '当前内容', '修改后内容（直接改写整行）']):
    table5b.rows[0].cells[i].text = h
    table5b.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

areas = [
    ('strategy',
     '数据战略咨询',
     '为企业量身定制数据战略规划与数字化转型路径，从顶层设计到实施路线图，确保每一步都有据可依。',
     '数据成熟度评估 / 转型路线图 / 组织与治理架构'),
    ('platform',
     '数据平台建设',
     '构建企业级数据平台，打通数据采集、存储、计算与分析全链路，让数据资产真正可用、可控、可运营。',
     '数据中台搭建 / 数据仓库建设 / 实时计算引擎'),
    ('governance',
     '数据治理服务',
     '建立完善的数据治理体系，保障数据质量、安全与合规，为企业数据资产保驾护航。',
     '数据标准制定 / 质量管理体系 / 安全合规保障'),
    ('ai',
     'AI应用落地',
     '将AI能力与企业场景深度融合，从智能决策到流程自动化，让前沿技术产生实实在在的业务价值。',
     '智能决策引擎 / 流程自动化 / 大模型应用'),
]
for i, (cid, title, desc, tags) in enumerate(areas):
    table5b.rows[i+1].cells[0].text = cid
    table5b.rows[i+1].cells[1].text = title
    table5b.rows[i+1].cells[2].text = f'{desc}\n\n标签：{tags}'
    table5b.rows[i+1].cells[3].text = ''

doc.add_paragraph()

# ════════════════════════════════════════════════════════
# 6. 使命·愿景·价值观
# ════════════════════════════════════════════════════════
doc.add_heading('六、「使命·愿景·价值观」区域', level=1)
doc.add_paragraph('数据文件：src/data/mission.json', style='Normal').italic = True

doc.add_heading('板块标题', level=2)
table6a = doc.add_table(rows=3, cols=3)
table6a.style = 'Light Grid Accent 1'
for i, h in enumerate(['字段', '当前内容', '修改后内容（留空=不改）']):
    table6a.rows[0].cells[i].text = h
    table6a.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

mission_title_items = [
    ('板块中文标题', '使命 · 愿景 · 价值观'),
    ('板块英文标题', 'MISSION & VISION & VALUES'),
]
for i, (label, value) in enumerate(mission_title_items):
    table6a.rows[i+1].cells[0].text = label
    table6a.rows[i+1].cells[1].text = value
    table6a.rows[i+1].cells[2].text = ''

doc.add_paragraph()

doc.add_heading('三列卡片', level=2)
table6b = doc.add_table(rows=4, cols=4)
table6b.style = 'Light Grid Accent 1'
for i, h in enumerate(['卡片', '中文标题 / 英文标题', '内容', '修改后内容（直接改写整行）']):
    table6b.rows[0].cells[i].text = h
    table6b.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

mission_items = [
    ('1', '使命 / MISSION', '让每个企业都能用数据驱动增长，释放数据要素的无限潜能，推动中国企业的数字化进程。'),
    ('2', '愿景 / VISION', '成为数据智能领域最受信赖的合作伙伴，打造中国企业数字化转型的标杆实践。'),
    ('3', '价值观 / VALUES', '务实求真——用数据说话，以结果为导向。创新进取——拥抱变化，持续探索前沿。协作共赢——与客户共同成长，以信任为基石。'),
]
for i, (num, title, content) in enumerate(mission_items):
    table6b.rows[i+1].cells[0].text = num
    table6b.rows[i+1].cells[1].text = title
    table6b.rows[i+1].cells[2].text = content
    table6b.rows[i+1].cells[3].text = ''

doc.add_paragraph()

# ════════════════════════════════════════════════════════
# 7. 新闻动态
# ════════════════════════════════════════════════════════
doc.add_heading('七、「新闻动态」区域', level=1)
doc.add_paragraph('数据文件：src/data/news.json', style='Normal').italic = True

doc.add_heading('板块标题', level=2)
table7a = doc.add_table(rows=3, cols=3)
table7a.style = 'Light Grid Accent 1'
for i, h in enumerate(['字段', '当前内容', '修改后内容（留空=不改）']):
    table7a.rows[0].cells[i].text = h
    table7a.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

news_title_items = [
    ('板块中文标题', '新闻动态'),
    ('板块英文标题', 'NEWS'),
]
for i, (label, value) in enumerate(news_title_items):
    table7a.rows[i+1].cells[0].text = label
    table7a.rows[i+1].cells[1].text = value
    table7a.rows[i+1].cells[2].text = ''

doc.add_paragraph()

doc.add_heading('新闻条目', level=2)
table7b = doc.add_table(rows=4, cols=4)
table7b.style = 'Light Grid Accent 1'
for i, h in enumerate(['条目', '日期 + 标题 + 摘要', '当前内容', '修改后内容（直接改写整行）']):
    table7b.rows[0].cells[i].text = h
    table7b.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

news_items = [
    ('1',
     '2025.06',
     '格数科技正式发布数据治理平台 V1.0',
     '经过团队数月的研发打磨，格数科技首款数据治理产品正式上线，标志着公司在数据智能化领域迈出坚实一步。'),
    ('2',
     '2025.05',
     '格数科技与行业伙伴达成战略合作',
     '携手多家行业领先企业，共同探索数据智能在垂直场景中的深度应用，拓展业务生态版图。'),
    ('3',
     '2025.04',
     '创始人受邀参加2025数据智能峰会',
     '格数科技创始人在峰会发表主题演讲，分享企业数字化转型的前沿思考与实践经验。'),
]
for i, (num, date, title, summary) in enumerate(news_items):
    table7b.rows[i+1].cells[0].text = num
    table7b.rows[i+1].cells[1].text = f'{date} | {title}'
    table7b.rows[i+1].cells[2].text = summary
    table7b.rows[i+1].cells[3].text = ''

doc.add_paragraph()

# ════════════════════════════════════════════════════════
# 8. 联系我们
# ════════════════════════════════════════════════════════
doc.add_heading('八、「联系我们」区域', level=1)
doc.add_paragraph('数据文件：src/data/contact.json + 组件内硬编码（ContactSection.jsx）', style='Normal').italic = True

doc.add_heading('板块标题 & 描述', level=2)
table8a = doc.add_table(rows=4, cols=3)
table8a.style = 'Light Grid Accent 1'
for i, h in enumerate(['字段', '当前内容', '修改后内容（留空=不改）']):
    table8a.rows[0].cells[i].text = h
    table8a.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

contact_title_items = [
    ('板块中文标题', '联系我们'),
    ('板块英文标题', 'CONTACT'),
    ('描述文字', '如果您对我们的解决方案感兴趣，欢迎随时与我们联系。我们期待与您一起探索数据智能的无限可能。'),
]
for i, (label, value) in enumerate(contact_title_items):
    table8a.rows[i+1].cells[0].text = label
    table8a.rows[i+1].cells[1].text = value
    table8a.rows[i+1].cells[2].text = ''

doc.add_paragraph()

doc.add_heading('联系信息 & UI文本（硬编码在 ContactSection.jsx 中）', level=2)
table8b = doc.add_table(rows=9, cols=3)
table8b.style = 'Light Grid Accent 1'
for i, h in enumerate(['字段', '当前内容', '修改后内容（留空=不改）']):
    table8b.rows[0].cells[i].text = h
    table8b.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

contact_ui_items = [
    ('邮箱标签', '电子邮箱'),
    ('邮箱地址', 'info@geshu-tech.com'),
    ('电话标签', '联系电话'),
    ('电话号', '（空）'),
    ('地址标签', '公司地址'),
    ('地址', '（空）'),
    ('邮件咨询按钮文本', '发送邮件咨询'),
    ('快速导航标题', '快速导航'),
]
for i, (label, value) in enumerate(contact_ui_items):
    table8b.rows[i+1].cells[0].text = label
    table8b.rows[i+1].cells[1].text = value
    table8b.rows[i+1].cells[2].text = ''

doc.add_paragraph()

doc.add_heading('快速导航链接', level=2)
table8c = doc.add_table(rows=5, cols=3)
table8c.style = 'Light Grid Accent 1'
for i, h in enumerate(['链接', '当前内容', '修改后内容（留空=不改）']):
    table8c.rows[0].cells[i].text = h
    table8c.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

quicklinks = [
    ('关于我们', '#about'),
    ('核心业务', '#business'),
    ('使命愿景', '#mission'),
    ('新闻动态', '#news'),
]
for i, (label, href) in enumerate(quicklinks):
    table8c.rows[i+1].cells[0].text = f'{label}（链接：{href}）'
    table8c.rows[i+1].cells[1].text = label
    table8c.rows[i+1].cells[2].text = ''

doc.add_paragraph()

# ════════════════════════════════════════════════════════
# 9. Footer
# ════════════════════════════════════════════════════════
doc.add_heading('九、页脚', level=1)
doc.add_paragraph('数据来源：Footer.jsx（动态年份 + site.json 数据）', style='Normal').italic = True

table9 = doc.add_table(rows=3, cols=3)
table9.style = 'Light Grid Accent 1'
for i, h in enumerate(['字段', '当前内容', '修改后内容（留空=不改）']):
    table9.rows[0].cells[i].text = h
    table9.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

footer_items = [
    ('版权文字', '© {year} Geshu Technology. 格数科技——企业级数据智能解决方案提供商。我们致力于帮助企业挖掘数据价值，实现数字化转型升级。'),
    ('ICP备案号', '（空，在 site.json 中为 icp 字段）'),
]
for i, (label, value) in enumerate(footer_items):
    table9.rows[i+1].cells[0].text = label
    table9.rows[i+1].cells[1].text = value
    table9.rows[i+1].cells[2].text = ''

doc.add_paragraph()

# ════════════════════════════════════════════════════════
# 10. HTML Meta 标签
# ════════════════════════════════════════════════════════
doc.add_heading('十、HTML Meta 标签 & SEO', level=1)
doc.add_paragraph('数据文件：index.html（头部 meta 标签）', style='Normal').italic = True

table10 = doc.add_table(rows=5, cols=3)
table10.style = 'Light Grid Accent 1'
for i, h in enumerate(['字段', '当前内容', '修改后内容（留空=不改）']):
    table10.rows[0].cells[i].text = h
    table10.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

meta_items = [
    ('页面标题 (title)', '格数科技 | 企业级数据智能解决方案'),
    ('Meta Description', '格数科技——企业级数据智能解决方案提供商。我们致力于让数据成为企业的核心驱动力。'),
    ('Meta Keywords', '格数科技,数据智能,数字化转型,数据战略,AI应用'),
    ('OG Title', '格数科技 | 企业级数据智能解决方案'),
]
for i, (label, value) in enumerate(meta_items):
    table10.rows[i+1].cells[0].text = label
    table10.rows[i+1].cells[1].text = value
    table10.rows[i+1].cells[2].text = ''

doc.add_paragraph()

# ════════════════════════════════════════════════════════
# 11. 其他硬编码 UI 文本
# ════════════════════════════════════════════════════════
doc.add_heading('十一、其他硬编码 UI 文本', level=1)
doc.add_paragraph('这些文本不在数据文件中，直接写在组件代码里。', style='Normal').italic = True

table11 = doc.add_table(rows=4, cols=4)
table11.style = 'Light Grid Accent 1'
for i, h in enumerate(['位置', '组件文件', '当前内容', '修改后内容']):
    table11.rows[0].cells[i].text = h
    table11.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

hardcoded_items = [
    ('Hero CTA按钮2', 'HeroSection.jsx → hero.json ctaSecondary', '联系我们'),
    ('新闻卡片「了解更多」链接', 'NewsCard.jsx', '了解更多'),
    ('关于我们区域空态', '—', '（无额外硬编码文本）'),
]
for i, (pos, file, value) in enumerate(hardcoded_items):
    table11.rows[i+1].cells[0].text = pos
    table11.rows[i+1].cells[1].text = file
    table11.rows[i+1].cells[2].text = value
    table11.rows[i+1].cells[3].text = ''

doc.add_paragraph()

# ── 页脚说明 ──
doc.add_paragraph()
doc.add_paragraph(
    '📝 使用说明：\n'
    '1. 在每个表格的「修改后内容」列填写新的文字内容；留空则表示不改动。\n'
    '2. 对于「直接改写整行」的单元格，如需修改，直接写新的标题+描述即可（用换行或"|"分隔）。\n'
    '3. 修改完成后将此 Word 文件发回给我，我会更新网站中对应的数据文件/组件代码。\n'
    '4. ⚠️ 请勿修改表格结构、合并单元格、或删除行。',
    style='Normal'
).italic = True

# ── 保存 ──
output_path = 'D:/hh/geshu-official-website/格数科技官网_内容清单.docx'
doc.save(output_path)
print(f'Done. Word doc saved to: {output_path}')
